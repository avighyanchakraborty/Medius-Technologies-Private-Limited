from flask import Flask, send_file, render_template_string
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

app = Flask(__name__)

HTML_PAGE = """
<!DOCTYPE html>
<html>
<head>
    <title>Document Generator</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            background: #f4f6f8;
            display: flex;
            height: 100vh;
            justify-content: center;
            align-items: center;
        }
        .box {
            background: black;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 4px 10px rgba(0,0,0,0.1);
            text-align: center;
        }
        button {
            padding: 12px 20px;
            font-size: 16px;
            background: #2563eb;
            color: white;
            border: none;
            border-radius: 6px;
            cursor: pointer;
        }
        button:hover {
            background: #1e40af;
        }
    </style>
</head>
<body>
    <div class="box">
        <h2>Generate Mediation Application</h2>
        <form action="/download" method="post">
            <button type="submit">Generate & Download DOCX</button>
        </form>
    </div>
</body>
</html>
"""

@app.route("/", methods=["GET"])
def home():
    return render_template_string(HTML_PAGE)

@app.route("/download", methods=["POST"])
def download_doc():
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    def center(text, size=12, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold

    def left(text, size=12, bold=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold

    # Document content
    center("FORM ‘A’", 14, True)
    center("MEDIATION APPLICATION FORM", 14, True)
    center("[REFER RULE 3(1)]")
    center("Mumbai District Legal Services Authority", 12, True)
    center("City Civil Court, Mumbai", 12, True)

    left("\nDETAILS OF PARTIES:")

    left("1. Name of Applicant")
    left("{{client_name}}\n")

    left("REGISTERED ADDRESS:")
    left("{{branch_address}}\n")

    left("CORRESPONDENCE BRANCH ADDRESS:")
    left("{{branch_address}}\n")

    left("Telephone No. {{mobile}}")
    left("Email ID info@kslegal.co.in\n")

    left("2. Name, Address and Contact details of Opposite Party:")
    left("Name {{customer_name}}\n")

    left("REGISTERED ADDRESS:")
    left("{{address1}} or __________________\n")

    left("CORRESPONDENCE ADDRESS:")
    left("{{address1}} or __________________\n")

    left("DETAILS OF DISPUTE:", bold=True)
    left("THE COMM. COURTS (PRE-INSTITUTION………SETTLEMENT) RULES, 2018")
    left(
        "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, "
        "2015 (4 of 2016):"
    )

    doc.save(temp.name)

    return send_file(
        temp.name,
        as_attachment=True,
        download_name="django_assignment.docx"
    )

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)
